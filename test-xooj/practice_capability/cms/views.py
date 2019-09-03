# -*-coding: utf-8 -*-
import os
import json
import time
import collections

from django.urls import reverse
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from django.db.models import ProtectedError
from django.shortcuts import render
from django.utils import timezone
from rest_framework import response, status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import exceptions
from event.utils.task import TaskHandler

from common_framework.utils.cache import CacheProduct, delete_cache
from common_framework.utils.rest.permission import IsStaffPermission
from common_framework.utils.image import save_image
from practice.api import get_task_object
from practice_capability import models as capability_models
from practice_capability.cms.viewset import TestPaperViewSet
from practice_capability.constant import PracticeResError
from practice_capability.models import TestPaper, TestPaperTask
from practice.models import TaskEvent
from practice_theory import models as theory_models
from course.models import Course, Lesson
from course import constant
from practice.api import PRACTICE_TYPE_THEORY
from practice.utils.task import get_type_by_hash, generate_task_hash
from practice_theory.models import ChoiceCategory



class Serializer:
    def __init__(self, raw, tpt):
        self.data = {
            'title': raw.title,
            'id': str(raw.id),
            'hash': str(raw.hash),
            'score': float(tpt.score),
            'content': raw.content
        }

        p_type = int(raw.hash.split('.')[-1])
        if p_type == 0:
            self.data['options'] = raw.option
            self.data['options_dsc'] = collections.OrderedDict(sorted(json.loads(raw.option).items(), key=lambda t: t[0]))
            self.data['is_choice_question'] = 1
            self.data['is_multiple_choice'] = 1 if raw.multiple else 0
        else:
            if (raw.url != None):
                self.data['url'] = raw.url
            if raw.file and raw.file.url:
                file_attach = {
                    'name': raw.file.name,
                    'url': raw.file.url,
                }

                self.data['file_url'] = file_attach



class SerializerNew:
    def __init__(self, raw, tpt):
        self.data = {
            'title': raw.title,
            'id': str(raw.id),
            'hash': str(raw.hash),
            'score': float(tpt.score),
            'content': raw.content
        }
        p_type = int(raw.hash.split('.')[-1])
        if p_type == 0:
            self.data['option'] = raw.option
            self.data['options_dsc'] = self.data['options_dsc'] = collections.OrderedDict(sorted(json.loads(raw.option).items(), key=lambda t: t[0]))
            self.data['multiple'] = raw.multiple


class LessonSerializer:
    def __init__(self, instance):
        self.data = {
            'lesson_id': instance.id,
            'lesson_name': instance.name
        }


def testpaper(request):
    return render(request, 'practice_capability/cms/testpapers.html')


def testpaper_detail(request, testpaper_id):
    new_type = request.GET.get('new_type', None)
    testpaper_id = int(testpaper_id)
    context = {
        'mode': 0,
    }
    if testpaper_id != 0:
        context = {
            'mode': 1,
        }
        testpaper = capability_models.TestPaper.objects.get(id=testpaper_id)

        context['testpaper_id'] = testpaper_id
        context['name'] = testpaper.name
        context['description'] = testpaper.introduction
        context['logo_url'] =testpaper.logo and testpaper.logo.url or ""
    else:
        context['description'] = ''
        context['name'] = ''
        context['new_type'] = new_type

        # 返回习题集数据
        event_list = TaskEvent.objects.exclude(status=0)
        course_list = Course.objects.filter(status=constant.COURSESTATE.NORMAL)
        # 返回理论基础分类

        choice_category_list = ChoiceCategory.objects.all()
        context["event_list"] = event_list
        context["course_list"] = course_list
        context["choice_category_list"] = choice_category_list

    context['testpaper_id'] = testpaper_id
    return render(request, 'practice_capability/cms/testpaper_detail_new.html', context)


@api_view(['GET', 'POST'])
def get_lesson_list(request):
    if request.method == 'POST':
        rows = []
        context = {}
        course_id = int(request.POST.get('course_id', None))
        lessons = Lesson.objects.filter(course__id=course_id)
        if lessons.exists():
            for lesson in lessons:
                rows.append(LessonSerializer(lesson).data)
        context['lesson_list'] = rows
        return Response(data=context, status=status.HTTP_200_OK)


def _judgment_question_type(hash):
    queryset = theory_models.ChoiceTask.objects.all()
    task = queryset.filter(hash=hash).first()
    if task:
        return task.multiple


@api_view(['GET', 'POST'])
def generate_docx(request):
    rows = []
    scores = []
    context = {}
    single_selections = []
    multiple_selections = []
    judgment_selections = []
    analysis_questions = []
    base_path = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    download_path = 'media/word/testpaper{}.docx'.format(int(time.time()))
    save_path = os.path.join(base_path, download_path)

    language = getattr(request, 'LANGUAGE_CODE', u'zh-hans')
    if language == 'en':
        JudgeTitle = 'Part 1:Judgment Problem'
        SingleChoice = 'Part 2:Single Choice'
        MultipleChoice = 'Part 3:Multiple Choice'
        ScoreCount = '{} tasks and {} PTS in total'
        Score = '({}PT for this task)'
    else:
        JudgeTitle = u'一、判断题'
        SingleChoice = u'二、单项选择题'
        MultipleChoice = u'三、多项选择题'
        ScoreCount = u'共{}赛题 满分{}PT'
        Score = u'（本题{}PT）'

    testpaper_id = int(request.GET.get("data"))
    testpaper = capability_models.TestPaper.objects.filter(id=testpaper_id).first()
    if testpaper:
        taskArrary = TestPaperTask.objects.filter(test_paper=testpaper)
        if taskArrary.exists():
            for t in taskArrary:
                task = get_task_object(t.task_hash)
                rows.append(Serializer(task, t).data)
            context['tasks'] = rows
            for task in context["tasks"]:
                if task.has_key("is_choice_question"):
                    scores.append(task['score'])
                    if _judgment_question_type(task["hash"]) == theory_models.ChoiceTask.TopicProblem.SINGLE:
                        single_selections.append(task)
                    elif _judgment_question_type(task["hash"]) == theory_models.ChoiceTask.TopicProblem.MULTIPLE:
                        multiple_selections.append(task)
                    else:
                        judgment_selections.append(task)
                else:
                    analysis_questions.append(task)

            context['name'] = testpaper.name
            context['allScore'] = sum(scores)
            context['number'] = len(single_selections) + len(multiple_selections) + len(judgment_selections)

            document = Document()
            document.styles['Normal'].font.name = u'宋体'
            document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

            p = document.add_heading(context['name'])
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            document.add_heading(ScoreCount.format(context['number'], context['allScore']), 4)

            document.add_heading(JudgeTitle, 2)
            for t in judgment_selections:
                document.add_paragraph(str(judgment_selections.index(t) + 1) + u'、' + t["content"])
                document.add_paragraph(Score.format(t['score']))
                for num, value in t["options_dsc"].items():
                    document.add_paragraph(u"{}、".format(num) + t["options_dsc"][num])

            document.add_heading(SingleChoice, 2)
            for t in single_selections:
                document.add_paragraph(str(single_selections.index(t)+1)+u'、'+t["content"])
                document.add_paragraph(Score.format(t['score']))
                for num, value in t["options_dsc"].items():
                    document.add_paragraph(u"{}、".format(num) + t["options_dsc"][num])

            document.add_heading(MultipleChoice, 2)
            for t in multiple_selections:
                document.add_paragraph(str(multiple_selections.index(t)+1)+u'、'+t["content"])
                document.add_paragraph(Score.format(t['score']))
                for num, value in t["options_dsc"].items():
                    document.add_paragraph(u"{}、".format(num) + t["options_dsc"][num])



            # document.add_heading(u'四、分析题', 2)
            # for t in analysis_questions:
            #     document.add_paragraph(str(analysis_questions.index(t) + 1) + u'、' + t.get("title", None))
            #     document.add_paragraph(u'（本题{}PT）'.format(t['score']))
            #     document.add_paragraph(t.get("content", None))

            # 保存文件
            if os.path.exists(os.path.join(base_path, 'media/word')):
                document.save(save_path)
            else:
                os.makedirs(os.path.join(base_path, 'media/word'))
                document.save(save_path)

            has_analysis = "yes" if analysis_questions else "no"

            return Response(data={"url": download_path, "has_analysis": has_analysis, "paper_name": testpaper.name},
                            status=status.HTTP_200_OK)


@api_view(['POST', 'GET'])
def ret_testpaper_detail(request, testpaper_id):
    if request.method == "GET":
        context = {}
        single_selections = []
        multiple_selections = []
        judgment_selections = []
        analysis_questions = []
        testpaper_id = int(testpaper_id)

        testpaper = capability_models.TestPaper.objects.get(id=testpaper_id)
        context['name'] = testpaper.name
        context['number'] = testpaper.task_number
        context['allScore'] = testpaper.task_all_score

        taskArrary = TestPaperTask.objects.filter(test_paper=testpaper)
        rows = []
        for t in taskArrary:
            task = get_task_object(t.task_hash)
            rows.append(Serializer(task, t).data)

        context['tasks'] = rows
        for task in context["tasks"]:
            if task.has_key("is_choice_question"):
                if _judgment_question_type(task["hash"]) == theory_models.ChoiceTask.TopicProblem.SINGLE:
                    single_selections.append(task)
                elif _judgment_question_type(task["hash"]) == theory_models.ChoiceTask.TopicProblem.MULTIPLE:
                    multiple_selections.append(task)
                else:
                    judgment_selections.append(task)
            else:
                analysis_questions.append(task)
        tasks = dict(
            judgment_selections=judgment_selections,
            single_selections=single_selections,
            multiple_selections=multiple_selections,
            analysis_questions=analysis_questions
        )
        context['tasks'] = tasks
        return response.Response({'error_code': 0, 'response_data': context})
    elif request.method == "POST":
        name = request.data['examname']
        if int(testpaper_id) == 0 and TestPaper.objects.filter(name=name).exists():
            return response.Response({'error_code': 1})

        questions = request.data['questions']
        list = questions.values()
        list.sort(key=lambda k: (k.get('qorder', 0)))
        number = 0
        allScore = 0

        for questions in list:
            number = number + 1
            allScore = allScore + int(questions['score'])

        event_teachers = request.data.get('teacher')

        if int(testpaper_id) > 0:
            testPaper = TestPaper.objects.get(id=testpaper_id)
            testPaper.name = name
            testPaper.task_number = number
            testPaper.task_all_score = allScore
            testPaper.teacher = ",".join(event_teachers) if event_teachers else None
            testPaper.save()

            tpt = TestPaperTask.objects.filter(test_paper=testPaper)
            for t in tpt:
                t.delete()

        else:
            testPaper = TestPaper.objects.create(
                name=name,
                task_number=number,
                task_all_score=allScore,
                create_time=timezone.now(),
                create_user=request.user,
                public=True,
            )

        for questions in list:
            TestPaperTask.objects.create(
                test_paper=testPaper,
                task_hash=questions['hash'],
                score=questions['score']
            )

        cache = CacheProduct("%s-%s" % (TestPaperViewSet.__module__, TestPaperViewSet.__name__))
        delete_cache(cache)

        return response.Response({'error_code': 1})


@api_view(['POST', "GET"])
@permission_classes((IsAuthenticated, IsStaffPermission,))
def handler_task_list(request, testpaper_id):
    examname = request.data.get('examname', '')
    examDescription = request.data.get('examDescription', '')
    datas = request.data.get('data', '')
    testpaper_id = int(testpaper_id)

    if request.method == "GET":
        rows = []
        if testpaper_id != 0:
            taskArrary = TestPaperTask.objects.filter(test_paper__id=testpaper_id)
            for t in taskArrary:
                task = get_task_object(t.task_hash)
                if not task:
                    continue
                data = SerializerNew(task, t).data
                rows.append(data)

        return Response(data=rows, status=status.HTTP_200_OK)

    if not datas or not json.loads(datas):
        raise exceptions.ParseError(PracticeResError.NO_EXAM_QUESTIONS)
    if not examname:
        raise exceptions.ParseError(PracticeResError.WARN_MESSAGES_7)
    if capability_models.TestPaper.objects.filter(status=capability_models.TestPaper.Status.NORMAL,name=examname).exclude(pk=testpaper_id).exists():
        raise exceptions.ParseError(PracticeResError.WARN_MESSAGES_6)

    datas = json.loads(datas)

    number = 0
    allScore = 0
    for data in datas:
        number = number + 1
        allScore = allScore + float(data['score'])

    logo_url = request.data.get("logo", None)

    if testpaper_id == 0:
        # 新增
        # 创建一张试卷
        if logo_url is not None and logo_url != '':
            logo = save_image(logo_url)
        else:
            logo = None
        testPaper = capability_models.TestPaper.objects.create(
            name=examname,
            task_number=number,
            task_all_score=allScore,
            introduction=examDescription,
            create_time=timezone.now(),
            create_user=request.user,
            logo=logo,
            public=True,
        )
    else:
        testPaper = capability_models.TestPaper.objects.filter(id=testpaper_id)
        if not testPaper:
            raise exceptions.NotFound(PracticeResError.NOT_FOUND_EXAM)
        testPaper = testPaper[0]
        testPaper.name = examname
        testPaper.task_number = number
        testPaper.task_all_score = allScore
        testPaper.introduction = examDescription
        if logo_url is not None and logo_url != '':
            logos = save_image(logo_url)
            testPaper.logo = logos
        testPaper.save()

        event_task_list = TestPaperTask.objects.filter(test_paper=testPaper)
        if event_task_list:
            # 判断原来是否存在初始数据, 批量删除
            try:
                event_task_list.delete()
            except ProtectedError:
                raise exceptions.ParseError(PracticeResError.CANNT_CHANGE_HAS_DONE)

    event_task_list =[]
    for data in datas:
        # copy_hash = copy_task(data['hash'])
        # if not copy_hash:
        #     continue

        hash_list = []
        if int(data['hash'].split('.')[-1]) != 0:
            hash_list.append(data['hash'])
            task_handler_class = TaskHandler
            copy_hash = task_handler_class.handle_tasks(hash_list)
            copy_hash = copy_hash[0]
        else:
            copy_hash = copy_task(data['hash'])
            if not copy_hash:
                continue

        event_task = TestPaperTask(
            test_paper=testPaper,
            task_hash=copy_hash,
            score=data['score']
        )
        event_task_list.append(event_task)
    TestPaperTask.objects.bulk_create(event_task_list)

    cache = CacheProduct("%s-%s" % (TestPaperViewSet.__module__, TestPaperViewSet.__name__))
    delete_cache(cache)

    return Response(data={'type': 'success'}, status=status.HTTP_200_OK)


def copy_task(taskhash):
    p_type = get_type_by_hash(taskhash)
    try:
        new_task = get_task_object(taskhash)
    except:
        return None

    new_task.pk = None
    new_task.is_copy = True
    new_task.hash = generate_task_hash(type=p_type)
    new_task.save()

    return new_task.hash


def paper_detail(request, testpaper_id):
    context = {
        'mode': 1
    }

    testpaper_id = int(testpaper_id)
    if testpaper_id == 0:
        context['mode'] = 0
    else:
        context['paper'] = TestPaper.objects.get(id=testpaper_id)

    return render(request, 'practice_capability/cms/paper_detail.html', context)


@api_view(['GET'])
@permission_classes((IsAuthenticated, IsStaffPermission,))
def share_teacher(request, testpaper_id):
    from common_cms import views as cms_views
    context = {
        'url_list_url': reverse("cms_practice_capability:testpaper"),
        'query_share_url': reverse("cms_practice_capability:api:test-paper-get-shares", kwargs={'pk': testpaper_id}),
        'modify_share_url': reverse("cms_practice_capability:api:test-paper-set-shares", kwargs={'pk': testpaper_id}),
    }

    return cms_views.share_teacher(request, context)

